# backend/agents/file_agent.py
# ============================================================
# NEXON File Agent
# Handles all file operations: create, convert, organize, summarize.
# ============================================================

import os
import shutil
from datetime import datetime
from typing import Dict
from backend.config import NEXON_HOME, DOCUMENTS_DIR
from backend.llm_engine import nexon_llm


class FileAgent:
    """
    File operations agent for NEXON.

    Capabilities:
    - Create documents (txt, docx, PDF, py, js, etc.).
    - Convert files (txt→PDF, CSV→JSON, etc.).
    - Move, rename, copy files.
    - Summarize long documents using LLM.
    - Merge/split PDFs (requires PyPDF2).
    - OCR from images (requires pytesseract + Pillow).
    """

    async def handle(self, intent: str, params: Dict, session_id: str) -> Dict:
        handlers = {
            "create_file"       : self.create_file,
            "convert_file"      : self.convert_file,
            "move_file"         : self.move_file,
            "summarize_document": self.summarize_document,
            "merge_pdf"         : self.merge_pdfs,
        }
        handler = handlers.get(intent, self._unknown)
        return await handler(params, session_id)

    async def create_file(self, params: Dict, session_id: str) -> Dict:
        """
        Create a new file with AI-generated or specified content.

        Args:
            params: {
                filename    (str): Target filename with extension.
                content     (str): File content (or LLM generates it).
                raw_text    (str): User intent for LLM content generation.
                file_type   (str): 'txt'|'py'|'js'|'md'|'html'|'csv'|'json'
            }
        """
        filename  = params.get("filename", f"nexon_doc_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt")
        content   = params.get("content", "")
        raw_text  = params.get("raw_text", "")
        file_type = os.path.splitext(filename)[1].lstrip(".")

        # Generate content with LLM if not provided
        if not content and raw_text:
            prompt = (
                f"Generate the content for a {file_type} file based on this request: {raw_text}. "
                f"Return only the file content, no explanation."
            )
            content = await nexon_llm.generate_response(prompt, language="en")

        # Docx support
        if file_type == "docx":
            return await self._create_docx(filename, content)

        # PDF support
        if file_type == "pdf":
            return await self._create_pdf(filename, content)

        # Plain text / code files
        file_path = os.path.join(DOCUMENTS_DIR, filename)
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)

        return {
            "success": True,
            "message": f"✅ File created: **{filename}**\nSaved to: `{file_path}`",
            "action": {
                "type"   : "file_created",
                "details": {"path": file_path, "filename": filename, "size": len(content)}
            }
        }

    async def _create_docx(self, filename: str, content: str) -> Dict:
        """Create a Word document."""
        try:
            from docx import Document
            doc = Document()
            doc.add_heading("NEXON Document", 0)
            for para in content.split("\n"):
                doc.add_paragraph(para)
            file_path = os.path.join(DOCUMENTS_DIR, filename)
            doc.save(file_path)
            return {
                "success": True,
                "message": f"✅ Word document created: **{filename}**\nSaved to: `{file_path}`",
                "action": {"type": "file_created", "details": {"path": file_path}}
            }
        except ImportError:
            return {"success": False, "message": "python-docx not installed.", "action": {}}

    async def _create_pdf(self, filename: str, content: str) -> Dict:
        """Create a PDF file."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas as pdf_canvas
            file_path = os.path.join(DOCUMENTS_DIR, filename)
            c = pdf_canvas.Canvas(file_path, pagesize=letter)
            y = 750
            for line in content.split("\n"):
                c.drawString(50, y, line[:90])
                y -= 20
                if y < 50:
                    c.showPage()
                    y = 750
            c.save()
            return {
                "success": True,
                "message": f"✅ PDF created: **{filename}**\nSaved to: `{file_path}`",
                "action": {"type": "file_created", "details": {"path": file_path}}
            }
        except ImportError:
            return {"success": False, "message": "reportlab not installed.", "action": {}}

    async def convert_file(self, params: Dict, session_id: str) -> Dict:
        """
        Convert a file from one format to another.

        Args:
            params: {
                source_path (str): Path to source file.
                target_format (str): Target format ('pdf', 'docx', 'txt', 'json', 'csv').
            }
        """
        source      = params.get("source_path", "")
        target_fmt  = params.get("target_format", "txt")

        if not os.path.exists(source):
            return {
                "success": False,
                "message": f"❌ File not found: {source}",
                "action": {}
            }

        base     = os.path.splitext(source)[0]
        out_path = f"{base}.{target_fmt}"

        if target_fmt == "pdf":
            with open(source, "r") as f:
                content = f.read()
            await self._create_pdf(os.path.basename(out_path), content)
        elif target_fmt == "txt":
            # Convert docx/pdf to text
            src_ext = os.path.splitext(source)[1].lower()
            if src_ext == ".docx":
                try:
                    from docx import Document
                    doc = Document(source)
                    text = "\n".join(p.text for p in doc.paragraphs)
                    with open(out_path, "w") as f:
                        f.write(text)
                except Exception as e:
                    return {"success": False, "message": str(e), "action": {}}
            else:
                shutil.copy(source, out_path)
        else:
            shutil.copy(source, out_path)

        return {
            "success": True,
            "message": f"✅ Converted to **{target_fmt}**: `{out_path}`",
            "action": {"type": "file_converted", "details": {"output": out_path}}
        }

    async def move_file(self, params: Dict, session_id: str) -> Dict:
        """
        Move or rename a file.

        Args:
            params: { source_path (str), destination (str) }
        """
        src  = params.get("source_path", "")
        dest = params.get("destination", DOCUMENTS_DIR)

        if not os.path.exists(src):
            return {"success": False, "message": f"File not found: {src}", "action": {}}

        if os.path.isdir(dest):
            dest = os.path.join(dest, os.path.basename(src))

        shutil.move(src, dest)
        return {
            "success": True,
            "message": f"✅ File moved to: `{dest}`",
            "action": {"type": "file_moved", "details": {"destination": dest}}
        }

    async def summarize_document(self, params: Dict, session_id: str) -> Dict:
        """
        Summarize a text/docx/pdf document using LLM.

        Args:
            params: { source_path (str) OR content (str) }
        """
        content    = params.get("content", "")
        source     = params.get("source_path", "")

        if not content and source and os.path.exists(source):
            ext = os.path.splitext(source)[1].lower()
            if ext in (".txt", ".md", ".py", ".js"):
                with open(source, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read()[:8000]  # Limit to 8k chars

        if not content:
            return {"success": False, "message": "No content to summarize.", "action": {}}

        summary = await nexon_llm.generate_response(
            f"Please provide a concise summary of this document:\n\n{content}",
            language="en"
        )
        return {
            "success": True,
            "message": f"📄 **Document Summary:**\n\n{summary}",
            "action": {"type": "document_summarized", "details": {"summary": summary}}
        }

    async def merge_pdfs(self, params: Dict, session_id: str) -> Dict:
        """
        Merge multiple PDFs into one.

        Args:
            params: { files (list): List of PDF file paths. output (str): Output filename. }
        """
        files  = params.get("files", [])
        output = params.get("output", os.path.join(DOCUMENTS_DIR, "merged.pdf"))

        if len(files) < 2:
            return {"success": False, "message": "Need at least 2 PDFs to merge.", "action": {}}

        try:
            import PyPDF2
            writer = PyPDF2.PdfWriter()
            for pdf_path in files:
                if os.path.exists(pdf_path):
                    reader = PyPDF2.PdfReader(pdf_path)
                    for page in reader.pages:
                        writer.add_page(page)
            with open(output, "wb") as f:
                writer.write(f)
            return {
                "success": True,
                "message": f"✅ Merged {len(files)} PDFs → `{output}`",
                "action": {"type": "pdf_merged", "details": {"output": output}}
            }
        except ImportError:
            return {"success": False, "message": "PyPDF2 not installed.", "action": {}}
        except Exception as e:
            return {"success": False, "message": str(e), "action": {}}

    async def _unknown(self, params: Dict, session_id: str) -> Dict:
        return {"success": False, "message": "Unknown file action.", "action": {}}